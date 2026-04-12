"""
AI Content Intelligence System — Text Extraction Module
Handles extraction from PDF, DOCX, and plain text files.
"""

import io


def extract_from_pdf(file_data):
    """Extract text from a PDF file."""
    import fitz  # PyMuPDF

    text_parts = []
    try:
        if isinstance(file_data, bytes):
            pdf = fitz.open(stream=file_data, filetype="pdf")
        else:
            pdf = fitz.open(stream=file_data.read(), filetype="pdf")

        for page in pdf:
            text_parts.append(page.get_text())
        pdf.close()
    except Exception as e:
        return f"[Error extracting PDF: {e}]"

    return "\n".join(text_parts)


def extract_from_docx(file_data):
    """Extract text from a DOCX file."""
    import docx

    try:
        if isinstance(file_data, bytes):
            doc = docx.Document(io.BytesIO(file_data))
        else:
            doc = docx.Document(file_data)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception as e:
        return f"[Error extracting DOCX: {e}]"


def extract_from_txt(file_data):
    """Extract text from a plain text file."""
    try:
        if isinstance(file_data, bytes):
            return file_data.decode("utf-8", errors="replace")
        else:
            return file_data.read().decode("utf-8", errors="replace")
    except Exception as e:
        return f"[Error reading text file: {e}]"


def extract_text(uploaded_file):
    """
    Main extraction dispatcher.
    Accepts a Streamlit UploadedFile object.
    """
    name = uploaded_file.name.lower()
    raw = uploaded_file.read()

    if name.endswith(".pdf"):
        return extract_from_pdf(raw)
    elif name.endswith(".docx"):
        return extract_from_docx(raw)
    elif name.endswith(".txt") or name.endswith(".md"):
        return extract_from_txt(raw)
    else:
        return extract_from_txt(raw)


def get_supported_formats():
    """Return list of supported file formats."""
    return ["pdf", "docx", "txt", "md"]
